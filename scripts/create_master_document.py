import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_master_markdown():
    """Combine all synthesis documents into a master markdown file"""

    reports_dir = 'C:/Projects/OSINT - Foresight/reports/country=IT'

    # Files to combine in order
    files_to_combine = [
        'ITALY_SYNTHESIS_ANALYSIS.md',
        'PHASE_01_SYNTHESIS.md',
        'PHASE_02_SYNTHESIS.md',
        'PHASE_03_SYNTHESIS.md',
        'PHASE_04_SYNTHESIS.md',
        'PHASE_05_SYNTHESIS.md',
        'PHASE_06_SYNTHESIS.md',
        'PHASE_07_SYNTHESIS.md',
        'PHASE_08_SYNTHESIS.md',
        'PHASE_09_SYNTHESIS.md',
        'PHASE_10_SYNTHESIS.md',
        'PHASE_11_SYNTHESIS.md',
        'PHASE_12_SYNTHESIS.md',
        'PHASE_13_SYNTHESIS.md'
    ]

    master_content = []

    # Add master document header
    master_content.append("# MASTER DOCUMENT: Italy Technology Security Assessment - Complete Analysis")
    master_content.append("")
    master_content.append(f"**Generated:** {datetime.now().strftime('%Y-%m-%d')}")
    master_content.append("**Document Type:** Comprehensive Synthesis")
    master_content.append("**Sources:** ChatGPT v6 Phase Analysis + Claude Enhancements")
    master_content.append("**Classification:** OSINT Synthesis")
    master_content.append("")
    master_content.append("---")
    master_content.append("")

    # Add table of contents
    master_content.append("## Table of Contents")
    master_content.append("")
    master_content.append("1. [Executive Synthesis Analysis](#executive-synthesis-analysis)")
    master_content.append("2. [Phase 1: Baseline & Setup](#phase-1-baseline--setup)")
    master_content.append("3. [Phase 2: Indicators & Data Sources](#phase-2-indicators--data-sources)")
    master_content.append("4. [Phase 3: Institutional Landscape](#phase-3-institutional-landscape)")
    master_content.append("5. [Phase 4: Supply Chain Analysis](#phase-4-supply-chain-analysis)")
    master_content.append("6. [Phase 5: Institutional Analysis](#phase-5-institutional-analysis)")
    master_content.append("7. [Phase 6: Funding Analysis](#phase-6-funding-analysis)")
    master_content.append("8. [Phase 7: International Links](#phase-7-international-links)")
    master_content.append("9. [Phase 8: Risk Assessment](#phase-8-risk-assessment)")
    master_content.append("10. [Phase 9: Governance & Posture](#phase-9-governance--posture)")
    master_content.append("11. [Phase 10: Early Warning Indicators](#phase-10-early-warning-indicators)")
    master_content.append("12. [Phase 11: Foresight & Adversarial Analysis](#phase-11-foresight--adversarial-analysis)")
    master_content.append("13. [Phase 12: Capacity Building](#phase-12-capacity-building)")
    master_content.append("14. [Phase 13: Implementation & Monitoring](#phase-13-implementation--monitoring)")
    master_content.append("")
    master_content.append("---")
    master_content.append("")

    # Process each file
    for i, filename in enumerate(files_to_combine):
        filepath = os.path.join(reports_dir, filename)

        if os.path.exists(filepath):
            print(f"Adding {filename}...")

            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()

            # Add section separator for main synthesis
            if i == 0:
                master_content.append("# PART I: EXECUTIVE SYNTHESIS ANALYSIS")
                master_content.append("")
            elif i == 1:
                master_content.append("---")
                master_content.append("")
                master_content.append("# PART II: DETAILED PHASE ANALYSIS")
                master_content.append("")

            # Add the content
            master_content.append(content)

            # Add page break marker for Word conversion
            if i < len(files_to_combine) - 1:
                master_content.append("")
                master_content.append("---")
                master_content.append("")
        else:
            print(f"Warning: {filename} not found")

    # Save master markdown
    master_md_path = os.path.join(reports_dir, 'ITALY_MASTER_SYNTHESIS.md')
    with open(master_md_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(master_content))

    print(f"Master markdown document created: {master_md_path}")
    return master_md_path

def parse_markdown_table(table_text):
    """Parse markdown table into rows and columns"""
    lines = table_text.strip().split('\n')
    if len(lines) < 3:
        return None

    # Parse header
    header = [cell.strip() for cell in lines[0].split('|')]
    header = [h for h in header if h]  # Remove empty strings

    # Parse data rows (skip separator line)
    data_rows = []
    for line in lines[2:]:
        row = [cell.strip() for cell in line.split('|')]
        row = [r for r in row if r != '']  # Keep cells, even if empty content
        if len(row) > 0:
            data_rows.append(row)

    return header, data_rows

def add_formatted_table(doc, table_text):
    """Add a properly formatted table to the Word document"""
    parsed = parse_markdown_table(table_text)
    if not parsed:
        return False

    header, data_rows = parsed
    if not header or not data_rows:
        return False

    # Create table
    table = doc.add_table(rows=1, cols=len(header))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Add header row
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(header):
        hdr_cells[i].text = text
        # Make header bold and centered
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add data rows
    for row_data in data_rows:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data[:len(header)]):
            row_cells[i].text = text

    return True

def convert_master_to_docx(markdown_path):
    """Convert master markdown to Word with enhanced formatting"""

    print("Converting master document to Word...")

    # Read markdown content
    with open(markdown_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create new document
    doc = Document()

    # Set document properties
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Process content line by line
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Handle page breaks
        if line.strip() == '---':
            # Check if this is a table separator
            if i > 0 and '|' in lines[i-1]:
                i += 1
                continue
            # Otherwise add page break
            doc.add_page_break()
            i += 1
            continue

        # Handle headers
        if line.startswith('#'):
            level = len(line.split(' ')[0])
            text = line.lstrip('#').strip()

            # Special formatting for main title
            if level == 1 and 'MASTER DOCUMENT' in text:
                heading = doc.add_heading(text, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 1 and 'PART' in text:
                heading = doc.add_heading(text, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Add some space after part headers
                doc.add_paragraph()
            elif level == 1:
                heading = doc.add_heading(text, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_heading(text, min(level-1, 4))

        # Handle metadata lines
        elif line.startswith('**') and ':' in line:
            p = doc.add_paragraph()
            # Parse the bold and regular parts
            parts = re.split(r'(\*\*[^*]+\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part:
                    p.add_run(part)

        # Handle bullet points
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # Check for bold text in bullet
            p = doc.add_paragraph(style='List Bullet')
            parts = re.split(r'(\*\*[^*]+\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part:
                    p.add_run(part)

        # Handle numbered lists
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s+', '', line.strip())
            doc.add_paragraph(text, style='List Number')

        # Handle tables
        elif '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            # Collect table lines
            table_lines = [line]
            j = i + 1
            while j < len(lines) and '|' in lines[j]:
                table_lines.append(lines[j])
                j += 1

            # Add table to document
            table_text = '\n'.join(table_lines)
            if not add_formatted_table(doc, table_text):
                # If table parsing fails, add as text
                for tline in table_lines:
                    doc.add_paragraph(tline, style='No Spacing')

            i = j - 1

        # Handle special formatting indicators
        elif line.strip().startswith('✓') or line.strip().startswith('✅'):
            p = doc.add_paragraph()
            p.add_run('✓ ').bold = True
            p.add_run(line.strip()[1:].strip())

        elif line.strip().startswith('❌') or line.strip().startswith('✗'):
            p = doc.add_paragraph()
            p.add_run('✗ ').bold = True
            p.add_run(line.strip()[1:].strip())

        # Regular paragraphs
        else:
            # Check for inline formatting
            p = doc.add_paragraph()

            # Handle bold text
            parts = re.split(r'(\*\*[^*]+\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part:
                    p.add_run(part)

        i += 1

    # Save the document
    output_path = markdown_path.replace('.md', '.docx')
    doc.save(output_path)
    print(f"Word document created: {output_path}")
    return output_path

def main():
    """Create master document in both markdown and Word formats"""

    print("Creating master synthesis document...")
    print("="*50)

    # Create master markdown
    master_md = create_master_markdown()

    print("")
    print("Converting to Word format with proper table formatting...")

    # Convert to Word
    master_docx = convert_master_to_docx(master_md)

    print("")
    print("="*50)
    print("Master documents created successfully!")
    print(f"  Markdown: {master_md}")
    print(f"  Word:     {master_docx}")

    return master_md, master_docx

if __name__ == '__main__':
    main()
