import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def parse_markdown_table(table_text):
    """Parse markdown table into rows and columns"""
    lines = table_text.strip().split('\n')
    if len(lines) < 3:  # Need at least header, separator, and one data row
        return None

    # Parse header
    header = [cell.strip() for cell in lines[0].split('|') if cell.strip()]

    # Skip separator line

    # Parse data rows
    data_rows = []
    for line in lines[2:]:
        row = [cell.strip() for cell in line.split('|') if cell.strip()]
        if row:
            data_rows.append(row)

    return header, data_rows

def convert_markdown_to_docx(markdown_file, output_file):
    """Convert a markdown file to Word document with proper formatting"""

    # Read markdown content
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Split content into lines
    lines = content.split('\n')

    i = 0
    while i < len(lines):
        line = lines[i]

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Handle headers
        if line.startswith('#'):
            level = len(line.split(' ')[0])
            text = line.lstrip('#').strip()
            if level == 1:
                heading = doc.add_heading(text, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_heading(text, level-1)

        # Handle metadata lines (bold:value format)
        elif '**' in line and ':' in line:
            p = doc.add_paragraph()
            parts = line.split('**')
            for j, part in enumerate(parts):
                if j % 2 == 1:  # Bold text
                    run = p.add_run(part)
                    run.bold = True
                else:
                    p.add_run(part)

        # Handle bullet points
        elif line.strip().startswith('-') or line.strip().startswith('*'):
            text = line.strip().lstrip('-*').strip()
            doc.add_paragraph(text, style='List Bullet')

        # Handle numbered lists
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s+', '', line.strip())
            doc.add_paragraph(text, style='List Number')

        # Check for tables
        elif '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            # Collect table lines
            table_lines = [line]
            j = i + 1
            while j < len(lines) and '|' in lines[j]:
                table_lines.append(lines[j])
                j += 1

            # Parse table
            parsed = parse_markdown_table('\n'.join(table_lines))
            if parsed:
                header, data_rows = parsed
                # Create table
                table = doc.add_table(rows=len(data_rows) + 1, cols=len(header))
                table.style = 'Table Grid'

                # Add header
                for col, text in enumerate(header):
                    cell = table.rows[0].cells[col]
                    cell.text = text
                    # Make header bold
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

                # Add data rows
                for row_idx, row_data in enumerate(data_rows):
                    for col_idx, text in enumerate(row_data[:len(header)]):
                        table.rows[row_idx + 1].cells[col_idx].text = text

            i = j - 1

        # Handle horizontal rules
        elif line.strip() == '---' or line.strip() == '***':
            doc.add_paragraph('_' * 50)

        # Regular paragraph
        else:
            # Check for inline formatting
            p = doc.add_paragraph()
            text = line

            # Handle bold text
            parts = re.split(r'(\*\*[^*]+\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        i += 1

    # Save document
    doc.save(output_file)
    return output_file

def main():
    """Convert all phase synthesis documents to Word format"""

    reports_dir = 'C:/Projects/OSINT - Foresight/reports/country=IT'

    # List of phase documents to convert
    phase_docs = [
        'ITALY_SYNTHESIS_ANALYSIS.md',
        'PHASE_01_SYNTHESIS.md',
        'PHASE_02_SYNTHESIS.md',
        'PHASE_03_SYNTHESIS.md',
        'PHASE_04_SYNTHESIS.md',
        'PHASE_05_SYNTHESIS.md',
        'PHASE_06_SYNTHESIS.md',
        'PHASE_07_SYNTHESIS.md',
        'PHASE_08_SYNTHESIS.md',
        'PHASE_09_SYNTHESIS.md',
        'PHASE_10_SYNTHESIS.md',
        'PHASE_11_SYNTHESIS.md',
        'PHASE_12_SYNTHESIS.md',
        'PHASE_13_SYNTHESIS.md'
    ]

    converted_files = []

    for doc_name in phase_docs:
        input_file = os.path.join(reports_dir, doc_name)
        output_file = os.path.join(reports_dir, doc_name.replace('.md', '.docx'))

        if os.path.exists(input_file):
            print(f"Converting {doc_name}...")
            try:
                # Skip if file already exists
                if os.path.exists(output_file):
                    print(f"  [SKIP] {os.path.basename(output_file)} already exists")
                    converted_files.append(output_file)
                else:
                    converted = convert_markdown_to_docx(input_file, output_file)
                    converted_files.append(converted)
                    print(f"  [OK] Created {os.path.basename(converted)}")
            except Exception as e:
                print(f"  [ERROR] Converting {doc_name}: {str(e)}")
        else:
            print(f"  [WARNING] File not found: {doc_name}")

    print(f"\nConversion complete! {len(converted_files)} documents converted to Word format.")
    print(f"Location: {reports_dir}")

    return converted_files

if __name__ == '__main__':
    main()
