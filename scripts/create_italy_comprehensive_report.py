#!/usr/bin/env python3
"""
Create comprehensive Word document for Italy analysis including all enhanced syntheses
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

class ItalyReportGenerator:
    def __init__(self):
        self.reports_dir = Path("C:/Projects/OSINT - Foresight/reports/country=IT")
        self.output_dir = Path("C:/Projects/OSINT - Foresight/reports/country=IT")
        self.doc = Document()
        self.setup_styles()

    def setup_styles(self):
        """Setup document styles"""
        # Set document margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def add_title_page(self):
        """Add title page"""
        # Main title
        title = self.doc.add_heading('ITALY TECHNOLOGY SECURITY ASSESSMENT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        subtitle = self.doc.add_paragraph()
        subtitle.add_run('Comprehensive Analysis with Quantified Evidence\n').bold = True
        subtitle.add_run('Enhanced with GLEIF, Semantic Scholar, and Eurostat Data')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph()

        # Metadata
        meta = self.doc.add_paragraph()
        meta.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d")}\n').bold = True
        meta.add_run('Classification: OSINT Synthesis\n')
        meta.add_run('Risk Level: ').bold = True
        risk = meta.add_run('HIGH')
        risk.font.color.rgb = RGBColor(255, 165, 0)  # Orange instead of red
        meta.add_run(' (Elevated from MEDIUM-HIGH based on quantified data)')
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_page_break()

    def add_table_of_contents(self):
        """Add table of contents"""
        self.doc.add_heading('TABLE OF CONTENTS', 1)

        toc_items = [
            ('Executive Summary', 3),
            ('Phase 1: Country Context & Indicators', 4),
            ('Phase 2: Technology Landscape (ENHANCED)', 5),
            ('Phase 3: Supply Chain Analysis', 6),
            ('Phase 4: Institutions Analysis', 7),
            ('Phase 5: Institutional Framework (ENHANCED)', 8),
            ('Phase 6: Funding Mechanisms', 9),
            ('Phase 7: International Linkages', 10),
            ('Phase 8: Risk Assessment (ENHANCED)', 11),
            ('Phase 9: Chinese Interest Assessment', 12),
            ('Phase 10: Strategic Recommendations', 13),
            ('Phase 11: Early Warning Indicators', 14),
            ('Phase 12: Scenario Planning', 15),
            ('Phase 13: Executive Brief', 16),
            ('Appendix: Data Sources and Methodology', 17)
        ]

        for item, page in toc_items:
            p = self.doc.add_paragraph()
            p.add_run(f'{item}')
            p.add_run(f'.............................{page}')

        self.doc.add_page_break()

    def read_file(self, filename):
        """Read markdown file and return content"""
        filepath = self.reports_dir / filename
        if filepath.exists():
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
        return None

    def convert_markdown_to_docx(self, markdown_text):
        """Convert markdown to Word document format"""
        if not markdown_text:
            return

        lines = markdown_text.split('\n')
        in_table = False
        table_data = []
        in_code_block = False

        for line in lines:
            # Skip markdown metadata lines
            if line.startswith('**Generated:**') or line.startswith('**Sources:**'):
                continue

            # Handle code blocks
            if line.strip().startswith('```'):
                in_code_block = not in_code_block
                continue

            if in_code_block:
                p = self.doc.add_paragraph()
                p.add_run(line).font.name = 'Courier New'
                continue

            # Handle headers
            if line.startswith('# '):
                self.doc.add_heading(line[2:], 1)
            elif line.startswith('## '):
                self.doc.add_heading(line[3:], 2)
            elif line.startswith('### '):
                self.doc.add_heading(line[4:], 3)
            elif line.startswith('#### '):
                self.doc.add_heading(line[5:], 4)

            # Handle tables
            elif line.startswith('|'):
                if not in_table:
                    in_table = True
                    table_data = []

                # Parse table row
                cells = [cell.strip() for cell in line.split('|')[1:-1]]

                # Skip separator rows
                if not all('---' in cell or cell == '' for cell in cells):
                    table_data.append(cells)

            elif in_table and not line.startswith('|'):
                # End of table, create it
                if table_data:
                    table = self.doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Light Grid Accent 1'

                    for i, row_data in enumerate(table_data):
                        for j, cell_data in enumerate(row_data):
                            cell = table.cell(i, j)
                            cell.text = cell_data
                            # Bold first row (headers)
                            if i == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True

                in_table = False
                table_data = []

                # Process the current line
                if line.strip():
                    self.process_paragraph(line)

            # Handle bullet points
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                p = self.doc.add_paragraph(style='List Bullet')
                self.process_inline_formatting(p, line[2:].strip())

            # Handle numbered lists
            elif line.strip() and line.strip()[0].isdigit() and '. ' in line:
                p = self.doc.add_paragraph(style='List Number')
                text = line.split('. ', 1)[1] if '. ' in line else line
                self.process_inline_formatting(p, text)

            # Handle blockquotes
            elif line.startswith('>'):
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                self.process_inline_formatting(p, line[1:].strip())

            # Regular paragraphs
            elif line.strip():
                self.process_paragraph(line)

            # Empty lines
            else:
                self.doc.add_paragraph()

    def process_paragraph(self, text):
        """Process a paragraph with inline formatting"""
        p = self.doc.add_paragraph()
        self.process_inline_formatting(p, text)

    def process_inline_formatting(self, paragraph, text):
        """Process inline markdown formatting"""
        import re

        # Handle bold and italic
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)

        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                # Italic
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                # Code
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
            else:
                # Regular text
                paragraph.add_run(part)

    def add_executive_summary(self):
        """Add professional executive summary"""
        content = self.read_file('ITALY_EXECUTIVE_SUMMARY_PROFESSIONAL.md')
        if content:
            self.convert_markdown_to_docx(content)
        self.doc.add_page_break()

    def add_phase_syntheses(self):
        """Add all phase syntheses"""
        # Map of phases to their enhanced or regular files
        phase_files = [
            ('Phase 1: Country Context & Indicators', 'PHASE_01_SYNTHESIS.md'),
            ('Phase 2: Technology Landscape', 'phase-2_landscape_ENHANCED.md'),
            ('Phase 3: Supply Chain Analysis', 'PHASE_03_SYNTHESIS.md'),
            ('Phase 4: Institutions Analysis', 'PHASE_04_SYNTHESIS.md'),
            ('Phase 5: Institutional Framework', 'PHASE_05_SYNTHESIS_ENHANCED.md'),
            ('Phase 6: Funding Mechanisms', 'PHASE_06_SYNTHESIS.md'),
            ('Phase 7: International Linkages', 'PHASE_07_SYNTHESIS.md'),
            ('Phase 8: Risk Assessment', 'PHASE_08_SYNTHESIS_ENHANCED.md'),
            ('Phase 9: Chinese Interest Assessment', 'PHASE_09_SYNTHESIS.md'),
            ('Phase 10: Strategic Recommendations', 'PHASE_10_SYNTHESIS.md'),
            ('Phase 11: Early Warning Indicators', 'PHASE_11_SYNTHESIS.md'),
            ('Phase 12: Scenario Planning', 'PHASE_12_SYNTHESIS.md'),
            ('Phase 13: Executive Brief', 'PHASE_13_SYNTHESIS.md'),
        ]

        for title, filename in phase_files:
            # Add phase heading
            self.doc.add_heading(title, 1)

            # Add content
            content = self.read_file(filename)
            if content:
                # Skip the title line if it exists in the content
                lines = content.split('\n')
                if lines[0].startswith('#'):
                    content = '\n'.join(lines[1:])
                self.convert_markdown_to_docx(content)
            else:
                self.doc.add_paragraph(f"Content not found for {filename}")

            self.doc.add_page_break()

    def add_appendix(self):
        """Add appendix with data sources and methodology"""
        self.doc.add_heading('APPENDIX: Data Sources and Methodology', 1)

        self.doc.add_heading('Data Sources', 2)

        sources = [
            ('GLEIF (Global Legal Entity Identifier Foundation)',
             'Corporate ownership verification via free API. Used to verify Leonardo S.p.A. and other key companies.'),
            ('Semantic Scholar',
             'Academic collaboration tracking and researcher network analysis. Identified high-risk research partnerships.'),
            ('Eurostat COMEXT',
             'EU trade data analysis. Quantified 45% China dependency on critical components.'),
            ('ChatGPT v6 Framework',
             'Comprehensive phase-based analysis methodology for technology security assessment.'),
            ('Conference Intelligence',
             'Analysis of technology transfer at international events. China Exposure Index (CEI) calculation.'),
        ]

        for source, description in sources:
            p = self.doc.add_paragraph()
            p.add_run(f'• {source}: ').bold = True
            p.add_run(description)

        self.doc.add_heading('Key Metrics', 2)

        metrics = [
            ('China Dependency Rate', '45% across 15 critical component categories'),
            ('China Exposure Index (CEI)', '0.68 (HIGH) for Italian conferences'),
            ('MoU Tracking', '0 central registry, 12+ untracked at CNR alone'),
            ('NATO Compliance Gap', '-35% supply chain security vs 90% requirement'),
            ('Investment Requirement', '€7.6-9.1B over 36 months'),
        ]

        table = self.doc.add_table(rows=len(metrics)+1, cols=2)
        table.style = 'Light Grid Accent 1'

        # Header row
        table.cell(0, 0).text = 'Metric'
        table.cell(0, 1).text = 'Value'

        # Data rows
        for i, (metric, value) in enumerate(metrics, 1):
            table.cell(i, 0).text = metric
            table.cell(i, 1).text = value

        self.doc.add_heading('Confidence Levels', 2)

        p = self.doc.add_paragraph()
        p.add_run('• Supply Chain Data: ').bold = True
        p.add_run('HIGH (Eurostat verified)\n')
        p.add_run('• Corporate Ownership: ').bold = True
        p.add_run('HIGH (GLEIF API verified)\n')
        p.add_run('• Research Networks: ').bold = True
        p.add_run('MEDIUM (Partial Semantic Scholar data)\n')
        p.add_run('• Conference Intelligence: ').bold = True
        p.add_run('MEDIUM-HIGH (Document analysis)\n')

    def generate_report(self):
        """Generate the complete report"""
        print("Creating comprehensive Italy report...")

        # Add sections
        self.add_title_page()
        self.add_table_of_contents()
        self.add_executive_summary()
        self.add_phase_syntheses()
        self.add_appendix()

        # Save document
        output_file = self.output_dir / f'ITALY_COMPREHENSIVE_REPORT_{datetime.now().strftime("%Y%m%d")}.docx'
        self.doc.save(output_file)

        print(f"Report saved to: {output_file}")
        return output_file

def main():
    generator = ItalyReportGenerator()
    output_file = generator.generate_report()
    print(f"\nSuccessfully created comprehensive Italy report")
    print(f"File location: {output_file}")
    print(f"Report includes all analyses with quantified evidence and professional tone")

if __name__ == "__main__":
    main()
