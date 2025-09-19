import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """
    Set cell borders.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        if border_name in kwargs:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), kwargs[border_name].get('val', 'single'))
            border.set(qn('w:sz'), str(kwargs[border_name].get('sz', 4)))
            border.set(qn('w:space'), str(kwargs[border_name].get('space', 0)))
            border.set(qn('w:color'), kwargs[border_name].get('color', '000000'))
            tcBorders.append(border)
    tcPr.append(tcBorders)

def convert_markdown_to_docx():
    # Create a new Document
    doc = Document()

    # Add title
    title = doc.add_heading('Italy Technology Security Assessment - Synthesis Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add metadata
    metadata = doc.add_paragraph()
    metadata.add_run('Generated: ').bold = True
    metadata.add_run('2025-09-16\n')
    metadata.add_run('Sources: ').bold = True
    metadata.add_run('ChatGPT v6 Phase Documents + Claude Analysis\n')
    metadata.add_run('Classification: ').bold = True
    metadata.add_run('OSINT Synthesis')

    # Key Correction section
    doc.add_heading('Key Correction', level=1)

    p = doc.add_paragraph(
        'The "9,278 Chinese STEM students" figure was erroneously included in initial analysis. '
        'ChatGPT\'s actual assessment is more nuanced and does not provide specific student numbers. '
        'The analysis focuses on:'
    )

    # Bullet points
    doc.add_paragraph('Visiting scholars at specific institutions under EU exchange programs', style='List Bullet')
    doc.add_paragraph('Talent circulation patterns without quantification', style='List Bullet')
    doc.add_paragraph('Security vetting gaps for inbound researchers', style='List Bullet')
    doc.add_paragraph('Outbound Italian PhDs taking postdocs in China', style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Critical Gap: ').bold = True
    p.add_run('No reliable data on actual numbers, fields of study, or institutional distribution of Chinese students/researchers in Italy.')

    # ChatGPT's Key Findings
    doc.add_heading("ChatGPT's Key Findings (from actual documents)", level=1)

    # 1. Conference Intelligence Framework
    doc.add_heading('1. Conference Intelligence Framework', level=2)
    doc.add_paragraph('ChatGPT emphasizes conference-enabled partnerships as primary vulnerability:')
    doc.add_paragraph('Tier-1 Events: Paris Air Show, Farnborough, DSEI, Eurosatory', style='List Bullet')
    doc.add_paragraph('Space/EO: IAC where ASI signed MoUs with CNSA-adjacent institutes (2020-2023)', style='List Bullet')
    doc.add_paragraph('Robotics: ICRA/IROS side meetings led to joint papers', style='List Bullet')
    doc.add_paragraph('Semiconductors: SEMICON Europa exploratory talks (undocumented outcomes)', style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('China Exposure Index (CEI)').bold = True
    p.add_run(' = china_presence_weighted × disclosure_risk × partnership_depth')
    doc.add_paragraph('Scaled 0-1 with 3.0 multiplier for Tier-1 critical events', style='List Bullet')
    doc.add_paragraph('Requires triad analysis (Italy-China-US co-presence)', style='List Bullet')

    # 2. Supply Chain Vulnerabilities
    doc.add_heading('2. Supply Chain Vulnerabilities', level=2)

    doc.add_heading('Critical Single Points of Failure', level=3)
    doc.add_paragraph('Harmonic Drive actuators (Germany) - global monopoly affecting robotics', style='List Bullet')
    doc.add_paragraph('GPU supply (NVIDIA/US) - no alternatives for HPC/quantum', style='List Bullet')
    doc.add_paragraph('Cryogenic dilution fridges (BlueFors/Finland, Oxford/UK) - quantum bottleneck', style='List Bullet')
    doc.add_paragraph('EUV lithography - complete ASML dependency', style='List Bullet')

    doc.add_heading('Dual-Exposure Risks', level=3)
    doc.add_paragraph('STMicroelectronics SiC production in both Catania and Shenzhen', style='List Bullet')
    doc.add_paragraph('Creates technology leakage pathway if Shenzhen operations compromise IP', style='List Bullet')

    # 3. Institutional Landscape
    doc.add_heading('3. Institutional Landscape', level=2)

    doc.add_heading('Research Centers with Confirmed China Engagement', level=3)
    doc.add_paragraph('CNR: 12 active MOUs with Chinese institutions, 47 ongoing projects', style='List Bullet')
    doc.add_paragraph('IIT (Italian Institute of Technology): Robotics/nanomaterials collaboration', style='List Bullet')
    doc.add_paragraph('Sant\'Anna School: PRC visiting scholars, joint robotics research', style='List Bullet')
    doc.add_paragraph('CINECA: Quantum panels with PRC teams at Q2B/EQTC', style='List Bullet')

    doc.add_heading('Critical Gaps in Oversight', level=3)
    doc.add_paragraph('No central MoU registry', style='List Bullet')
    doc.add_paragraph('Weak vetting of visiting scholars', style='List Bullet')
    doc.add_paragraph('Spin-out acquisitions not screened for PRC limited partners', style='List Bullet')
    doc.add_paragraph('Academic funding from PRC often hidden in bilateral projects', style='List Bullet')

    # 4. Technology Transfer Mechanisms - Table
    doc.add_heading('4. Technology Transfer Mechanisms', level=2)
    doc.add_paragraph('ChatGPT identifies four primary pathways:')

    # Create table for Technology Transfer Mechanisms
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Pathway'
    hdr_cells[1].text = 'Volume'
    hdr_cells[2].text = 'Control'
    hdr_cells[3].text = 'Mechanism'

    # Make header row bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Data rows
    data = [
        ['Academic Collaboration', 'HIGH', 'WEAK', 'Joint research, shared labs'],
        ['Industrial Partnerships', 'MEDIUM-HIGH', 'MODERATE', 'JVs, licensing, M&A'],
        ['Talent Recruitment', 'UNKNOWN (estimated significant)', 'MINIMAL', 'No tracking of Thousand Talents participation'],
        ['Conference Networking', 'HIGH', 'NONE', 'Initial contact point for partnerships']
    ]

    for i, row_data in enumerate(data, 1):
        row_cells = table.rows[i].cells
        for j, text in enumerate(row_data):
            row_cells[j].text = text

    # 5. Funding Landscape Risks
    doc.add_heading('5. Funding Landscape Risks', level=2)
    doc.add_paragraph('PNRR: €15-20bn for digital/tech with weak foreign subcontractor vetting', style='List Bullet')
    doc.add_paragraph('Horizon Europe: PRC entities as "Swiss substitutes" in consortia', style='List Bullet')
    doc.add_paragraph('Regional Funds: Lombardia/Lazio innovation funds lack ownership transparency', style='List Bullet')
    doc.add_paragraph('Venture Capital: PRC LPs in EU funds creating indirect access', style='List Bullet')

    # 6. Early Warning Indicators
    doc.add_heading('6. Early Warning Indicators (EWI)', level=2)
    doc.add_paragraph('ChatGPT proposes specific metrics:')
    doc.add_paragraph('TED Procurement: 90-day rolling counts on security-relevant CPVs', style='List Bullet')
    doc.add_paragraph('Standards Momentum: Quarterly role weights (member=1, rapporteur=3, editor=5)', style='List Bullet')
    doc.add_paragraph('DIANA/NATO Pipeline: Additions to Italy-linked accelerators/test centers', style='List Bullet')
    doc.add_paragraph('Conference Metrics: Attendance rates, repeaters, triad co-appearances', style='List Bullet')

    # 7. Negative Evidence
    doc.add_heading('7. Negative Evidence (What\'s NOT Happening)', level=2)
    doc.add_paragraph('ChatGPT explicitly notes absence of:')
    doc.add_paragraph('Direct PRC involvement in Italian nuclear sector', style='List Bullet')
    doc.add_paragraph('Confirmed military technology transfers', style='List Bullet')
    doc.add_paragraph('Large-scale Chinese ownership in aerospace primes', style='List Bullet')
    doc.add_paragraph('Systematic targeting of Arctic technologies (Italy not Arctic-adjacent)', style='List Bullet')

    # Refined Intelligence Requirements
    doc.add_heading('Refined Intelligence Requirements', level=1)

    doc.add_heading('Priority Data Gaps', level=2)

    doc.add_heading('1. Quantified Academic Presence', level=3)
    doc.add_paragraph('Actual numbers of Chinese researchers by institution and field', style='List Bullet')
    doc.add_paragraph('Breakdown: undergrad vs graduate vs postdoc vs visiting', style='List Bullet')
    doc.add_paragraph('Temporal trends accounting for COVID impact (2020-2022 baseline disruption)', style='List Bullet')
    doc.add_paragraph('Geographic distribution (NUTS-2 regions)', style='List Bullet')

    doc.add_heading('2. Conference Intelligence', level=3)
    doc.add_paragraph('Complete rosters for Tier-1/2 events (2020-2024)', style='List Bullet')
    doc.add_paragraph('Side meeting records and MoU timings', style='List Bullet')
    doc.add_paragraph('Technology disclosure assessments at panel/session level', style='List Bullet')

    doc.add_heading('3. Supply Chain Mapping', level=3)
    doc.add_paragraph('Complete CAGE/NCAGE registry with China exposure flags', style='List Bullet')
    doc.add_paragraph('Tier-2/3 supplier dependencies', style='List Bullet')
    doc.add_paragraph('Replacement feasibility assessments', style='List Bullet')

    doc.add_heading('4. Funding Transparency', level=3)
    doc.add_paragraph('LEI parent chains for all grant recipients', style='List Bullet')
    doc.add_paragraph('Ultimate beneficial ownership for spin-outs', style='List Bullet')
    doc.add_paragraph('Hidden PRC funding through third countries', style='List Bullet')

    # Nuanced Risk Assessment
    doc.add_heading('Nuanced Risk Assessment', level=1)

    doc.add_heading('High-Confidence Risks', level=2)
    doc.add_paragraph('Conference-enabled partnerships creating unvetted technology transfer', style='List Bullet')
    doc.add_paragraph('Supply chain single points of failure (actuators, GPUs, cryogenics)', style='List Bullet')
    doc.add_paragraph('Weak oversight of visiting scholars in dual-use domains', style='List Bullet')
    doc.add_paragraph('Opaque funding chains enabling indirect PRC access', style='List Bullet')

    doc.add_heading('Medium-Confidence Concerns', level=2)
    doc.add_paragraph('Scale of talent recruitment programs (evidence of activity, not quantified)', style='List Bullet')
    doc.add_paragraph('Technology leakage through joint ventures', style='List Bullet')
    doc.add_paragraph('Standards influence through committee positions', style='List Bullet')
    doc.add_paragraph('Regional innovation fund exploitation', style='List Bullet')

    doc.add_heading('Low-Confidence/Speculative', level=2)
    doc.add_paragraph('Specific numbers of Chinese students/researchers', style='List Bullet')
    doc.add_paragraph('Direct military technology transfer', style='List Bullet')
    doc.add_paragraph('Systematic IP theft (activity yes, systematic orchestration uncertain)', style='List Bullet')

    # Recommendations for Next Phase
    doc.add_heading('Recommendations for Next Phase', level=1)

    # Immediate Actions Table
    doc.add_heading('Immediate Actions', level=2)
    table2 = doc.add_table(rows=5, cols=2)
    table2.style = 'Table Grid'

    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = 'Action #'
    hdr_cells2[1].text = 'Description'

    for cell in hdr_cells2:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    actions = [
        ['1', 'Deploy conference monitoring for upcoming 2025 events'],
        ['2', 'Map supply chain dependencies with criticality scoring'],
        ['3', 'Establish MoU registry requirement for all research institutions'],
        ['4', 'Implement visiting scholar vetting protocols']
    ]

    for i, action in enumerate(actions, 1):
        row_cells = table2.rows[i].cells
        row_cells[0].text = action[0]
        row_cells[1].text = action[1]

    # Data Collection Priorities
    doc.add_heading('Data Collection Priorities', level=2)
    doc.add_paragraph('TED Database Mining: Extract all procurement awards to Chinese entities', style='List Bullet')
    doc.add_paragraph('Conference Archives: Reconstruct 2020-2024 attendance patterns', style='List Bullet')
    doc.add_paragraph('Patent Analysis: Map co-assignments and technology flows', style='List Bullet')
    doc.add_paragraph('Academic Census: Quantify foreign researcher presence properly', style='List Bullet')

    doc.add_heading('Analytical Focus', level=2)
    doc.add_paragraph('Move beyond "China = bad" to specific vulnerability identification', style='List Bullet')
    doc.add_paragraph('Map legitimate collaboration vs exploitation pathways', style='List Bullet')
    doc.add_paragraph('Account for COVID baseline disruptions in trend analysis', style='List Bullet')
    doc.add_paragraph('Distinguish between presence and actual risk', style='List Bullet')

    # Quality Assessment of ChatGPT Analysis
    doc.add_heading('Quality Assessment of ChatGPT Analysis', level=1)

    doc.add_heading('Strengths', level=2)
    doc.add_paragraph('Comprehensive framework covering 13 phases', style='List Bullet')
    doc.add_paragraph('Strong emphasis on conference intelligence (often overlooked)', style='List Bullet')
    doc.add_paragraph('Supply chain vulnerability identification', style='List Bullet')
    doc.add_paragraph('Clear data gap acknowledgment', style='List Bullet')

    doc.add_heading('Weaknesses', level=2)
    doc.add_paragraph('Lack of quantitative baselines', style='List Bullet')
    doc.add_paragraph('Limited specific examples/cases', style='List Bullet')
    doc.add_paragraph('Minimal use of open-source intelligence', style='List Bullet')
    doc.add_paragraph('Arctic analysis included despite Italy\'s non-Arctic status', style='List Bullet')

    doc.add_heading('Missing Elements', level=2)
    doc.add_paragraph('No analysis of Belt and Road legacy impacts (Italy withdrew 2024)', style='List Bullet')
    doc.add_paragraph('Limited coverage of port infrastructure (Genoa, Trieste critical)', style='List Bullet')
    doc.add_paragraph('Insufficient focus on maritime domain', style='List Bullet')
    doc.add_paragraph('No discussion of EU regulatory environment impact', style='List Bullet')

    # Conclusion
    doc.add_heading('Conclusion', level=1)

    doc.add_paragraph(
        'ChatGPT\'s Italy analysis provides a robust framework but lacks specific intelligence data. '
        'The emphasis on conference-enabled partnerships and supply chain vulnerabilities is valuable. '
        'However, critical gaps remain in quantifying academic presence, tracking funding flows, '
        'and mapping actual technology transfer incidents.'
    )

    doc.add_paragraph('\nThe erroneous "9,278 Chinese STEM students" figure highlighted the importance of:')
    doc.add_paragraph('Verifying specific numbers with primary sources', style='List Bullet')
    doc.add_paragraph('Avoiding reductive "China = bad" narratives', style='List Bullet')
    doc.add_paragraph('Distinguishing between presence and actual risk', style='List Bullet')
    doc.add_paragraph('Accounting for temporal factors (COVID disruption)', style='List Bullet')

    doc.add_paragraph(
        '\nNext steps should focus on populating ChatGPT\'s framework with actual intelligence data '
        'from TED procurement records, conference archives, and patent databases.'
    )

    # Save the document
    doc.save('C:/Projects/OSINT - Foresight/reports/country=IT/ITALY_SYNTHESIS_ANALYSIS.docx')
    print('Word document created successfully: ITALY_SYNTHESIS_ANALYSIS.docx')

if __name__ == '__main__':
    convert_markdown_to_docx()
