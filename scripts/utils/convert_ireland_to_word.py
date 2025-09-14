"""
Convert Ireland OSINT Foresight reports from Markdown to Word document
with proper table formatting
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

def parse_markdown_table(table_text):
    """Parse a markdown table into rows and columns"""
    lines = table_text.strip().split('\n')
    rows = []

    for line in lines:
        # Skip separator lines (containing only dashes and pipes)
        if re.match(r'^[\s\-\|]+$', line):
            continue
        # Split by pipe and clean up
        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        if cells:
            rows.append(cells)

    return rows

def add_markdown_section(doc, content, level=1):
    """Add a markdown section to the document, converting tables and formatting"""
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        # Handle headers
        if line.startswith('#'):
            header_level = len(line) - len(line.lstrip('#'))
            header_text = line.lstrip('#').strip()
            if header_level == 1:
                heading = doc.add_heading(header_text, level=1)
            elif header_level == 2:
                heading = doc.add_heading(header_text, level=2)
            elif header_level == 3:
                heading = doc.add_heading(header_text, level=3)
            else:
                p = doc.add_paragraph()
                run = p.add_run(header_text)
                run.bold = True
            i += 1

        # Detect and handle tables
        elif '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            # Start of a table
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1

            # Parse the table
            table_data = parse_markdown_table('\n'.join(table_lines))

            if table_data:
                # Create Word table
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Light Grid Accent 1'

                # Populate table
                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_data in enumerate(row_data):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = cell_data
                        # Bold first row (header)
                        if row_idx == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True

            # Add spacing after table
            doc.add_paragraph()

        # Handle bullet points
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(line[2:])
            i += 1

        # Handle numbered lists
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style='List Number')
            p.add_run(re.sub(r'^\d+\.\s', '', line))
            i += 1

        # Handle bold text
        elif line:
            p = doc.add_paragraph()
            # Process inline formatting
            text = line

            # Handle bold text (**text** or __text__)
            parts = re.split(r'(\*\*[^*]+\*\*|__[^_]+__)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('__') and part.endswith('__'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
            i += 1

        else:
            # Empty line
            i += 1

def main():
    # Create a new Word document
    doc = Document()

    # Set document title
    title = doc.add_heading('Ireland OSINT Foresight Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Comprehensive Intelligence Assessment')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    doc.add_page_break()

    # Define the order of phases
    phase_files = [
        ('phase-0_taxonomy.md', 'Phase 0: Definitions and Taxonomy'),
        ('phase-1_setup.md', 'Phase 1: Setup and Configuration'),
        ('phase-2_indicators.md', 'Phase 2: Indicators and Data Sources'),
        ('phase-3_landscape.md', 'Phase 3: Technology Landscape'),
        ('phase-4_supply_chain.md', 'Phase 4: Supply Chain Security'),
        ('phase-5_institutions.md', 'Phase 5: Institutional Mapping'),
        ('phase-6_funders.md', 'Phase 6: Funding Analysis'),
        ('phase-7_links.md', 'Phase 7: International Links'),
        ('phase-8_risk.md', 'Phase 8: Risk Assessment'),
        ('phase-9_posture.md', 'Phase 9: PRC Interest & MCF'),
        ('phase-10_redteam.md', 'Phase 10: Red Team Analysis'),
        ('phase-11_foresight.md', 'Phase 11: Strategic Foresight')
    ]

    reports_dir = Path('reports/country=IE')

    for filename, phase_title in phase_files:
        file_path = reports_dir / filename

        if file_path.exists():
            print(f"Processing {filename}...")

            # Add phase as main heading
            doc.add_heading(phase_title, 1)

            # Read and process the markdown content
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # Add the content to the document
            add_markdown_section(doc, content, level=2)

            # Add page break between major sections
            if filename != phase_files[-1][0]:  # Don't add page break after last section
                doc.add_page_break()
        else:
            print(f"Warning: {filename} not found")

    # Save the document
    output_path = 'Ireland_OSINT_Foresight_Analysis.docx'
    doc.save(output_path)
    print(f"\nWord document created: {output_path}")
    print(f"Location: {os.path.abspath(output_path)}")

if __name__ == '__main__':
    main()
